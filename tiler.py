#!/usr/bin/env python
# coding: utf8
'''
Simple image tiler

Collects images from the current folder and below and
tiles them into specified MS Word document. 

    Usage: tiler.py <filename> <image_width>

        <filename> -- MS Word document (filename.docx will be created)
        <image_width>  -- width of image (in cm) inserted to the word document; 
                          all images (inside the docx) will have the same width.

    The script looks for images with names of the form: somename_xx.ext, where:
        ext -- one of the allowed extensions (*.jpg, *.bmp, *.png, *.jpeg);
        xx -- the number of repetition of the image in the docx.
    
    Image files don't match this criterion will be ignored (not inserted into docx). 
    
    Note: 
        You need to prepend `./` to `tiler.py` if working in the directory 
        not included in PATH env-variable.
    
Author: Dmitry E. Kislov
e-mail: kislov@easydan.com
Date: 2016, May 18
'''

from __future__ import print_function
from docx import Document
from docx.shared import Inches
from docx.section import Section
from docx.enum.section import WD_SECTION
import os, sys, re


# ------------- New docx settings -------------
document = Document()
section = document.sections[0]
section.start_type = WD_SECTION.NEW_PAGE
# ----------- Page settings -------------------
section.page_height = Inches(11.69)
section.page_width = Inches(8.27)
section.left_margin = Inches(0.1)
section.right_margin = Inches(0.1)
section.top_margin = Inches(0.1)
section.bottom_margin = Inches(0.1)
# ----------------------------------------------

numpat = re.compile(r'.+_(\d+)\..+')

if len(sys.argv)!=3:
    print(__doc__)
    sys.exit(0)

filename = sys.argv[1]
if os.path.exists(filename):
    print('File <{0}> already exists. Provide another filename.'.format(filename))
    sys.exit(1)


imsize = sys.argv[2]
try:
    imsize = float(imsize)
except:
    print('You need to specify <imsize> parameter (in cm).')
    print(__doc__)
    sys.exit(1)
    
# ---------- Paragraph formatting --------------
paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph.left_indent = Inches(0.0)
paragraph.right_indent = Inches(0.0)
paragraph.first_line_indent = Inches(0.0)
paragraph.space_before = Inches(0.0)
paragraph.space_after = Inches(0.0)
paragraph.style = document.styles['No Spacing']
# ----------------------------------------------

run = paragraph.add_run()

for d, dirs, files in os.walk('/home/dmitry/Documents/EXCHANGE/ЭТИКЕТКИ'):
    count = 0
    for f in files:
        if os.path.splitext(f)[1].lower() in ['.png', '.jpg', '.jpeg', '.bmp']:
            if numpat.match(f):
                count += 1
                try:
                    thenum = int(numpat.findall(f)[0])
                    for j in range(thenum):
                        run.add_picture(os.path.join(d.decode('utf8'), f.decode('utf8')), width=Inches(imsize/2.54))
                except ValueError:
                    pass
    try:
        os.remove(os.path.join(d, os.path.basename(d) + '_' + filename + '.docx'))
    except (IOError, OSError):
        pass

    try:
        if count > 0:
            document.save(os.path.join(d, os.path.basename(d) + '_' + filename + '.docx'))
    except (IOError, OSError):
        pass

